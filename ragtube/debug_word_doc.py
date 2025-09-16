#!/usr/bin/env python3
"""Debug Word document structure."""

from docx import Document

def debug_word_document(doc_path: str):
    """Debug the structure of the Word document."""
    try:
        doc = Document(doc_path)
        print(f"📖 Document has {len(doc.paragraphs)} paragraphs")
        
        for i, paragraph in enumerate(doc.paragraphs[:20]):  # First 20 paragraphs
            text = paragraph.text.strip()
            if text:
                print(f"Para {i}: {text[:100]}...")
        
        print("\n" + "="*50)
        print("Looking for video patterns...")
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if "VIDEO" in text.upper() or "====" in text:
                print(f"Found pattern at para {i}: {text}")
        
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    debug_word_document("../CCBC Transcripts.docx")
