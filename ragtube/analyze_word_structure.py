#!/usr/bin/env python3
"""
Analyze the structure of the updated CCBC Transcripts.docx file.
"""

import os
import re
from docx import Document
from typing import List, Dict, Any

def analyze_word_structure(doc_path: str):
    """Analyze the structure of the Word document."""
    print(f"📖 Analyzing Word document: {doc_path}")
    
    try:
        doc = Document(doc_path)
        print(f"📄 Total paragraphs: {len(doc.paragraphs)}")
        
        videos = []
        current_video = None
        video_count = 0
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Check for video separator patterns
            if text.startswith("Video:") or text.startswith("VIDEO:") or text.startswith("video:"):
                video_count += 1
                video_id = text.replace("Video:", "").replace("VIDEO:", "").replace("video:", "").strip()
                print(f"\n📹 Video #{video_count}: {video_id}")
                
                if current_video:
                    videos.append(current_video)
                
                current_video = {
                    "video_id": video_id,
                    "title": "",
                    "transcript": "",
                    "paragraph_count": 0
                }
            
            elif text.startswith("Title:") or text.startswith("TITLE:") or text.startswith("title:"):
                if current_video:
                    title = text.replace("Title:", "").replace("TITLE:", "").replace("title:", "").strip()
                    current_video["title"] = title
                    print(f"   📝 Title: {title}")
            
            elif current_video and text:
                current_video["transcript"] += text + "\n"
                current_video["paragraph_count"] += 1
                
                # Show first few lines of transcript
                if current_video["paragraph_count"] <= 3:
                    print(f"   📄 Line {current_video['paragraph_count']}: {text[:100]}...")
        
        # Add last video
        if current_video:
            videos.append(current_video)
        
        print(f"\n🎯 Analysis Summary:")
        print(f"   📊 Total videos found: {len(videos)}")
        
        for i, video in enumerate(videos, 1):
            transcript_length = len(video["transcript"])
            word_count = len(video["transcript"].split())
            print(f"   📹 Video {i}: {video['video_id']}")
            print(f"      📝 Title: {video['title']}")
            print(f"      📄 Paragraphs: {video['paragraph_count']}")
            print(f"      📏 Characters: {transcript_length}")
            print(f"      🔤 Words: {word_count}")
            print(f"      📝 Preview: {video['transcript'][:200]}...")
            print()
        
        return videos
        
    except Exception as e:
        print(f"❌ Error analyzing document: {e}")
        return []

def main():
    """Main function."""
    doc_path = "../CCBC Transcripts.docx"
    
    if not os.path.exists(doc_path):
        print(f"❌ Word document not found: {doc_path}")
        return
    
    videos = analyze_word_structure(doc_path)
    
    if videos:
        print(f"✅ Successfully analyzed {len(videos)} videos!")
    else:
        print("❌ No videos found in document")

if __name__ == "__main__":
    main()
